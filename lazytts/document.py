"""Extract text — and chapter structure — from supported document formats."""
from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from . import textnorm


@dataclass
class Chapter:
    title: str
    text: str


def extract_chapters(path: str | Path, *, normalize: bool = True,
                     expand: bool = False, lang: str = "en") -> list[Chapter]:
    """Return the document split into chapters. Always at least one chapter.

    normalize -> de-hyphenate + strip page numbers.
    expand    -> also expand numbers & abbreviations to spoken words (in *lang*).
    """
    path = Path(path)
    ext = path.suffix.lower()
    if ext == ".txt":
        raw = _txt_chapters(_from_txt(path))
    elif ext == ".pdf":
        raw = _pdf_chapters(path)
    elif ext == ".epub":
        raw = _epub_chapters(path)
    elif ext == ".docx":
        raw = _docx_chapters(path)
    else:
        raise ValueError(
            f"Unsupported file type '{ext}'. Supported: .txt .pdf .epub .docx"
        )

    chapters: list[Chapter] = []
    for i, ch in enumerate(raw, 1):
        text = ch.text
        if normalize:
            text = textnorm.normalize(text, expand=expand, lang=lang)
        text = _clean(text)
        if not text:
            continue
        title = ch.title.strip() or f"Chapter {i}"
        chapters.append(Chapter(title=title, text=text))

    if not chapters:
        raise ValueError("No readable text found in the document.")
    return chapters


def extract_text(path: str | Path) -> str:
    """Whole-document text (chapters joined). Kept for convenience."""
    return "\n\n".join(c.text for c in extract_chapters(path))


def extract_metadata(path: str | Path) -> dict:
    """Best-effort {title, author} from document metadata; title falls back to
    the file name."""
    path = Path(path)
    ext = path.suffix.lower()
    title, author = path.stem, ""
    try:
        if ext == ".epub":
            from ebooklib import epub

            book = epub.read_epub(str(path))
            t = book.get_metadata("DC", "title")
            a = book.get_metadata("DC", "creator")
            if t and t[0] and t[0][0]:
                title = t[0][0]
            if a and a[0] and a[0][0]:
                author = a[0][0]
        elif ext == ".pdf":
            import fitz

            with fitz.open(path) as doc:
                meta = doc.metadata or {}
            title = (meta.get("title") or "").strip() or title
            author = (meta.get("author") or "").strip()
        elif ext == ".docx":
            import docx

            cp = docx.Document(str(path)).core_properties
            title = (cp.title or "").strip() or title
            author = (cp.author or "").strip()
    except Exception:
        pass
    return {"title": (title or path.stem).strip(), "author": (author or "").strip()}


# ── per-format extractors ────────────────────────────────────────

def _from_txt(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="latin-1")


# A short line that looks like a chapter/part heading (English + German).
_HEADING = re.compile(
    r"^\s*(chapter|part|book|section|kapitel|teil|buch|abschnitt|prolog|epilog)\b.*$"
    r"|^\s*\d{1,3}\.?\s+\S.*$",
    re.IGNORECASE,
)


def _txt_chapters(text: str) -> list[Chapter]:
    chapters: list[Chapter] = []
    cur_title: str | None = None
    cur: list[str] = []

    def flush():
        if cur and any(line.strip() for line in cur):
            chapters.append(Chapter(cur_title or "", "\n".join(cur)))

    for line in text.splitlines():
        stripped = line.strip()
        if stripped and len(stripped) < 60 and _HEADING.match(stripped):
            flush()
            cur_title = stripped
            cur = []
        else:
            cur.append(line)
    flush()
    return chapters or [Chapter("", text)]


def _pdf_chapters(path: Path) -> list[Chapter]:
    import fitz  # PyMuPDF

    with fitz.open(path) as doc:
        page_count = doc.page_count
        pages = [doc[i].get_text("text") for i in range(page_count)]
        toc = doc.get_toc(simple=True)  # [[level, title, page(1-based)], ...]

    if not toc:
        return [Chapter("", "\n".join(pages))]

    # Prefer top-level bookmarks; fall back to the whole TOC.
    tops = [(t[1], t[2]) for t in toc if t[0] == 1] or [(t[1], t[2]) for t in toc]

    chapters: list[Chapter] = []
    for idx, (title, start) in enumerate(tops):
        end = tops[idx + 1][1] if idx + 1 < len(tops) else page_count + 1
        s = max(start, 1) - 1
        e = min(end, page_count + 1) - 1
        chapters.append(Chapter(title, "\n".join(pages[s:e])))
    return chapters


def _epub_chapters(path: Path) -> list[Chapter]:
    """Split an EPUB by its table-of-contents sections (what a reader shows),
    not by internal spine files (which are often per-page fragments)."""
    import ebooklib
    from bs4 import BeautifulSoup, NavigableString, Tag
    from ebooklib import epub

    book = epub.read_epub(str(path))
    docs = {it.get_name(): it for it in book.get_items_of_type(ebooklib.ITEM_DOCUMENT)}
    by_base: dict[str, object] = {}
    for name, it in docs.items():
        by_base.setdefault(name.split("/")[-1], it)

    def resolve(fname: str):
        return docs.get(fname) or by_base.get(fname.split("/")[-1])

    # Flatten the TOC into ordered (title, file, anchor).
    flat: list[tuple[str, str, str | None]] = []

    def walk(entries):
        for e in entries:
            if isinstance(e, (tuple, list)):
                sec = e[0]
                children = e[1] if len(e) > 1 else []
                href = getattr(sec, "href", "") or ""
                title = getattr(sec, "title", "") or ""
                if href:
                    f, _, a = href.partition("#")
                    flat.append((title, f, a or None))
                walk(children)
            else:
                href = getattr(e, "href", "") or ""
                title = getattr(e, "title", "") or ""
                if href:
                    f, _, a = href.partition("#")
                    flat.append((title, f, a or None))

    walk(book.toc)
    if not flat:
        return _epub_spine_chapters(book, docs)

    # Group TOC entries by target file, preserving order.
    per_file: dict[str, list[tuple[str, str | None]]] = {}
    order: list[str] = []
    for title, f, a in flat:
        if f not in per_file:
            per_file[f] = []
            order.append(f)
        per_file[f].append((title, a))

    chapters: list[Chapter] = []
    for f in order:
        item = resolve(f)
        if item is None:
            continue
        soup = BeautifulSoup(item.get_content(), "html.parser")
        for tag in soup(["script", "style"]):
            tag.decompose()
        entries = per_file[f]
        anchors = {a: t for t, a in entries if a}

        if len(entries) <= 1 or not anchors:
            chapters.append(Chapter(entries[0][0], soup.get_text(separator="\n")))
            continue

        # Multiple TOC anchors in one file -> split at those anchors.
        body = soup.body or soup
        current = entries[0][0]
        buf: dict[str, list[str]] = {current: []}
        seq = [current]
        for node in body.descendants:
            if isinstance(node, Tag):
                nid = node.get("id")
                if nid and nid in anchors:
                    current = anchors[nid]
                    if current not in buf:
                        buf[current] = []
                        seq.append(current)
            elif isinstance(node, NavigableString):
                if str(node).strip():
                    buf[current].append(str(node))
        for t in seq:
            chapters.append(Chapter(t, " ".join(buf[t])))

    return chapters or _epub_spine_chapters(book, docs)


def _epub_spine_chapters(book, docs) -> list[Chapter]:
    """Fallback: one chapter per spine document (used only if there's no TOC)."""
    from bs4 import BeautifulSoup

    ordered = []
    for entry in book.spine:
        idref = entry[0] if isinstance(entry, (tuple, list)) else entry
        item = next((it for it in docs.values() if getattr(it, "id", None) == idref), None)
        if item is not None:
            ordered.append(item)
    if not ordered:
        ordered = list(docs.values())

    chapters: list[Chapter] = []
    for item in ordered:
        soup = BeautifulSoup(item.get_content(), "html.parser")
        for tag in soup(["script", "style"]):
            tag.decompose()
        heading = soup.find(["h1", "h2", "h3"])
        title = heading.get_text(" ", strip=True) if heading else ""
        chapters.append(Chapter(title, soup.get_text(separator="\n")))
    return chapters


def _docx_chapters(path: Path) -> list[Chapter]:
    import docx  # python-docx

    document = docx.Document(str(path))
    chapters: list[Chapter] = []
    cur_title: str | None = None
    cur: list[str] = []

    def flush():
        if cur and any(line.strip() for line in cur):
            chapters.append(Chapter(cur_title or "", "\n".join(cur)))

    for para in document.paragraphs:
        style = (para.style.name or "") if para.style else ""
        if style.startswith("Heading 1") or style == "Title":
            flush()
            cur_title = para.text.strip()
            cur = []
        else:
            cur.append(para.text)
    flush()
    return chapters or [Chapter("", "\n".join(p.text for p in document.paragraphs))]


def _clean(text: str) -> str:
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r" *\n *", "\n", text)
    return text.strip()
